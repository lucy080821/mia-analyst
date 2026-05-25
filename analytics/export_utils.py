import io
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from xhtml2pdf import pisa
import re
import base64
import io

# --- Font Registration for PDF (Unicode support) ---
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Use local font to support Vietnamese and avoid permission issues
FONT_DIR = os.path.join(os.path.dirname(__file__), "fonts")
FONT_PATH = os.path.join(FONT_DIR, "arial.ttf")
IS_FONT_REGISTERED = False

def register_fonts():
    """Đăng ký phông chữ Arial hỗ trợ Unicode cho ReportLab."""
    global IS_FONT_REGISTERED
    if IS_FONT_REGISTERED:
        return True
    
    try:
        if os.path.exists(FONT_PATH):
            pdfmetrics.registerFont(TTFont('Arial', FONT_PATH))
            IS_FONT_REGISTERED = True
            return True
    except Exception as e:
        print(f"DEBUG: Font registration failed: {e}")
    return False

def sanitize_for_pdf(text):
    """
    Xóa hoặc thay thế Emoji và ký tự đặc biệt không được hỗ trợ bởi PDF font chuẩn.
    Hầu hết các phông chữ PDF (như Arial) không hỗ trợ Emoji (🐍, 📈, etc).
    """
    if not text:
        return ""
    
    # Một số emoji phổ biến hay gặp trong insight
    replacements = {
        "🐍": "[Python]",
        "📈": "[Forecast]",
        "📉": "[Trend Down]",
        "🔍": "[Insight]",
        "📊": "[Chart]",
        "⚠️": "[Lưu ý]",
        "✅": "[OK]",
        "❌": "[X]",
        "💡": "[Gợi ý]",
        "🤖": "[Mia AI]",
        "📋": "[Mục lục]",
        "✨": "[Highlight]"
    }
    
    for emoji, label in replacements.items():
        text = text.replace(emoji, label)
        
    # Loại bỏ các ký tự non-BMP (Emojis còn lại)
    # xhtml2pdf/ReportLab thường crash hoặc hiện ô vuông với các ký tự này.
    return "".join(c for c in text if ord(c) < 0xFFFF)

def markdown_to_html(text):
    """Chuyển đổi markdown đơn giản thành HTML cho PDF."""
    # Xử lý list
    lines = text.split('\n')
    in_list = False
    new_lines = []
    for line in lines:
        if line.strip().startswith('- '):
            if not in_list:
                new_lines.append('<ul>')
                in_list = True
            new_lines.append(f"<li>{line.strip()[2:]}</li>")
        else:
            if in_list:
                new_lines.append('</ul>')
                in_list = False
            new_lines.append(line)
    if in_list:
        new_lines.append('</ul>')
    
    text = '\n'.join(new_lines)
    
    # Xử lý các tag khác
    text = re.sub(r'### (.*)', r'<h3>\1</h3>', text)
    text = re.sub(r'## (.*)', r'<h2>\1</h2>', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)
    
    # Handle Images: ![Alt](data:image/...) -> <img src="data:image/...">
    text = re.sub(r'!\[.*?\]\((data:image\/.*?;base64,.*?)\)', r'<img src="\1" style="max-width: 100%; height: auto; display: block; margin: 15px auto;">', text)

    # Thay thế xuống dòng (tránh nháy kép trong list tag)
    # Chỉ replace \n thành <br> nếu không nằm trong tag html
    parts = re.split(r'(<.*?>)', text)
    for i in range(len(parts)):
        if not parts[i].startswith('<'):
            parts[i] = parts[i].replace('\n', '<br>')
    
    return ''.join(parts)

def generate_word_report(title, content):
    """Tạo file Word từ nội dung báo cáo, hỗ trợ cả Markdown và HTML Table."""
    import lxml.html
    doc = Document()
    
    # Set default font to Arial
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Title
    header = doc.add_heading(title, 0)
    for run in header.runs:
        run.font.name = 'Arial'
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # We'll treat the content as a series of blocks.
    # Tables are identified by <table tags.
    # Text blocks are split by tables.
    
    # Pre-process content: handle markdown images first
    # This keeps them on their own lines for easier detection
    content = content.replace('![Chart]', '\n![Chart]')
    
    # Split into chunks of [text, table, text, table...]
    chunks = re.split(r'(<table.*?</table>)', content, flags=re.DOTALL)
    
    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk: continue
        
        if chunk.startswith('<table'):
            try:
                # Use fragment_fromstring for better fragment handling
                table_el = lxml.html.fragment_fromstring(chunk)
                html_rows = table_el.xpath('.//tr')
                if not html_rows: continue
                
                # Determine max columns
                max_cols = 0
                for r in html_rows:
                    max_cols = max(max_cols, len(r.xpath('./td | ./th')))
                
                if max_cols == 0: continue
                
                table = doc.add_table(rows=len(html_rows), cols=max_cols)
                table.style = 'Table Grid'
                
                is_metric = 'metric-grid' in chunk
                
                for i, row_el in enumerate(html_rows):
                    cells_el = row_el.xpath('./td | ./th')
                    for j, cell_el in enumerate(cells_el):
                        if j >= max_cols: break
                        cell = table.cell(i, j)
                        
                        if is_metric:
                            # Handle metric-card structure
                            labels = cell_el.xpath('.//div[contains(@class, "metric-label")]/text()')
                            values = cell_el.xpath('.//div[contains(@class, "metric-value")]/text()')
                            if labels and values:
                                cell.text = f"{labels[0].strip()}\n{values[0].strip()}"
                            else:
                                cell.text = cell_el.text_content().strip()
                        else:
                            cell.text = cell_el.text_content().strip()
                        
                        # Style cell font
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Arial'
                                run.font.size = Pt(9 if is_metric else 10)
            except Exception as e:
                p = doc.add_paragraph(f"[Table Content Error: {str(e)}]")
                p.runs[0].font.name = 'Arial'
        else:
            # Handle Text Chunk
            lines = chunk.split('\n')
            for line in lines:
                line = line.strip()
                if not line: continue
                
                if line.startswith('## '):
                    h = doc.add_heading(line.replace('## ', '').strip(), level=1)
                    for run in h.runs:
                        run.font.name = 'Arial'
                        run.font.color.rgb = RGBColor(107, 70, 193)
                elif line.startswith('### '):
                    h = doc.add_heading(line.replace('### ', '').strip(), level=2)
                    for run in h.runs:
                        run.font.name = 'Arial'
                        run.font.color.rgb = RGBColor(30, 41, 59)
                elif line.startswith('![Chart](data:image'):
                    try:
                        b64_data = line.split('base64,')[1].replace(')', '')
                        image_data = base64.b64decode(b64_data)
                        doc.add_picture(io.BytesIO(image_data), width=Inches(5.5))
                    except: pass
                elif line.startswith('- '):
                    p = doc.add_paragraph(style='List Bullet')
                    _add_styled_text(p, line[2:].strip())
                else:
                    p = doc.add_paragraph()
                    _add_styled_text(p, line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def _add_styled_text(paragraph, text):
    """Helper to add bold text to a paragraph while stripping leftover HTML."""
    # Strip any stray HTML tags that might be in the text
    text = re.sub(r'<[^>]*>', '', text)
    
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part.replace('**', ''))
            run.bold = True
        else:
            run = paragraph.add_run(part)
        run.font.name = 'Arial'

def generate_pdf_report(title, content):
    """Tạo file PDF từ nội dung báo cáo hỗ trợ Tiếng Việt."""
    try:
        # Register font once
        print("DEBUG: Registering fonts...")
        register_fonts()
        
        # Sanitize content for PDF
        print("DEBUG: Sanitizing content...")
        clean_title = sanitize_for_pdf(title)
        clean_content = sanitize_for_pdf(content)
        
        # Prep font path for CSS (avoid backslashes in f-string and use file URI)
        css_font_path = "file:///" + FONT_PATH.replace('\\', '/')
        
        print("DEBUG: Preparing HTML content...")
        html_content = f"""
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @font-face {{
                    font-family: Arial;
                    src: url("{css_font_path}");
                }}
                @page {{ size: a4; margin: 2cm; }}
                body {{ font-family: Arial, sans-serif; color: #334155; line-height: 1.5; font-size: 11pt; }}
                h1 {{ font-family: Arial; color: #6b46c1; text-align: center; border-bottom: 2px solid #6b46c1; padding-bottom: 10px; font-size: 24pt; font-weight: bold; }}
                h2 {{ font-family: Arial; color: #6b46c1; margin-top: 20px; font-size: 18pt; border-bottom: 1px solid #e2e8f0; font-weight: bold; }}
                h3 {{ font-family: Arial; color: #1e293b; margin-top: 15px; font-size: 14pt; font-weight: bold; }}
                ul {{ margin-left: 20px; }}
                li {{ margin-bottom: 5px; font-family: Arial; }}
                p {{ font-family: Arial; }}
                strong {{ color: #0f172a; font-weight: bold; font-family: Arial; }}
                table {{ width: 100%; border-collapse: collapse; margin: 15px 0; font-size: 10pt; }}
                th {{ background-color: #f8fafc; color: #1e293b; font-weight: bold; border: 1px solid #e2e8f0; padding: 8px; text-align: left; }}
                td {{ border: 1px solid #e2e8f0; padding: 8px; color: #475569; }}
                .metric-grid {{ width: 100%; margin: 20px 0; }}
                .metric-card {{ 
                    width: 48%; 
                    background-color: #f8fafc; 
                    border: 1px solid #e2e8f0; 
                    padding: 15px; 
                    text-align: center; 
                    vertical-align: middle;
                }}
                .metric-label {{ 
                    font-size: 8pt; 
                    color: #64748b; 
                    font-weight: bold; 
                    text-transform: uppercase; 
                    margin-bottom: 5px; 
                }}
                .metric-value {{ 
                    font-size: 14pt; 
                    color: #1e293b; 
                    font-weight: bold; 
                }}
                .footer {{ position: fixed; bottom: 0; width: 100%; text-align: center; font-size: 9pt; color: #94a3b8; font-family: Arial; }}
            </style>
        </head>
        <body>
            <h1>{clean_title}</h1>
            {markdown_to_html(clean_content)}
            <div class="footer">Báo cáo được tạo bởi Mia Assistant</div>
        </body>
        </html>
        """
        buffer = io.BytesIO()
        # Explicitly set encoding to utf-8
        print("DEBUG: Calling pisa.CreatePDF...")
        pisa_status = pisa.CreatePDF(html_content, dest=buffer, encoding='utf-8')
        if pisa_status.err:
            print(f"DEBUG: Pisa error: {pisa_status.err}")
            return None
        buffer.seek(0)
        print("DEBUG: PDF generated successfully.")
        return buffer
    except Exception as e:
        print(f"DEBUG: generate_pdf_report exception: {str(e)}")
        return None
